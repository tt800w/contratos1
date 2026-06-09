import sys
import docx
from docx.shared import Pt

adult_docx_paths = [
    'public/contratos/Condiciones Específicas-Recursos Propios Mayor de Edad.docx',
    'public/contratos/Condiciones Específicas-Financiación Lumni- Mayor de Edad.docx',
    'public/contratos/Condiciones Específicas-Pronto Pago Mayor de Edad.docx'
]

carta_text = [
    ("Señores,", False),
    ("CAMPUSLANDS S.A.S. BIC", False),
    ("NIT. 901.628.406-1", False),
    ("", False),
    ("REFERENCIA: CARTA DE AUTORIZACIÓN DE USO DE IMAGEN, VOZ Y CESIÓN GRATUITA DE DERECHOS PATRIMONIALES", True),
    ("", False),
    ("________________, identificado(a) con cédula de ciudadanía No. ___________________ obrando en nombre propio, con domicilio en __________________, quien en adelante se denominará el CAMPER, por medio del presente documento autorizo de manera previa, expresa, informada y gratuita a CAMPUSLANDS S.A.S. BIC para captar, grabar, reproducir, editar, publicar, comunicar, divulgar y utilizar mi imagen, voz, nombre, testimonios y demás participaciones que se generen con ocasión de mi proceso de admisión, formación, actividades académicas, institucionales, promocionales, comerciales o de empleabilidad, en cualquier medio físico o digital, incluyendo redes sociales, páginas web, piezas publicitarias, material audiovisual, fotográfico, institucional y pedagógico.", False),
    ("", False),
    ("Igualmente, autorizo a CAMPUSLANDS para usar los contenidos que yo genere o aporte voluntariamente para tales fines y, en caso de que dichos contenidos constituyan obras protegidas por derechos de autor y hayan sido creados específicamente para CAMPUSLANDS con ocasión de mi vinculación al programa, cedo a título gratuito los derechos patrimoniales que legalmente sean transferibles, para su uso, reproducción, transformación, comunicación pública y demás formas de explotación permitidas por la ley, sin perjuicio de mis derechos morales como autor(a).", False),
    ("", False),
    ("Declaro que esta autorización y cesión se otorgan sin que haya lugar a pago, honorario, regalía, comisión o remuneración adicional alguna a mi favor.", False),
    ("", False),
    ("Así mismo, manifiesto que conozco y acepto que el tratamiento de mis datos personales se realizará conforme a la Política de Tratamiento de Datos Personales de CAMPUSLANDS.", False),
    ("", False),
    ("El CAMPER,", True),
    ("", False),
    ("", False),
    ("_____________________________", False),
    ("{NOMBRE DEL CAMPER}", True),
    ("C.C. No. {NUMERO DE CEDULA}", True),
    ("{CORREO}", True)
]

for doc_path in adult_docx_paths:
    print(f"Processing {doc_path}...")
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"Error opening {doc_path}: {e}")
        continue

    pagare_found = False
    senores_count = 0
    
    # Apply page break to Pagare and Senores
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith("PAGARÉ No.") or text.startswith("PAGARE No."):
            p.paragraph_format.page_break_before = True
            pagare_found = True
        elif text == "Señores," or text == "Señores":
            senores_count += 1
            if senores_count == 1:
                p.paragraph_format.page_break_before = True

    # Append Carta de Autorización
    doc.add_page_break()
    for text, bold in carta_text:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        if text:
            run = p.add_run(text)
            if bold:
                run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

    doc.save(doc_path)
    print(f"Saved {doc_path}")

print("Done.")
